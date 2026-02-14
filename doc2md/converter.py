"""
Document conversion functions
"""

import io
import re
from typing import Optional

import fitz  # pymupdf
from docx import Document
from markdownify import markdownify as md
from bs4 import BeautifulSoup
import httpx


# markdown.new as fallback service
MARKDOWN_NEW_URL = "https://markdown.new"


def convert_pdf(content: bytes, extract_images: bool = False) -> str:
    """
    Convert PDF to Markdown.
    
    Args:
        content: PDF file bytes
        extract_images: Whether to extract and embed images (base64)
    
    Returns:
        Markdown text
    """
    doc = fitz.open(stream=content, filetype="pdf")
    markdown_parts = []
    
    for page_num, page in enumerate(doc, 1):
        # Extract text blocks
        text = page.get_text("text")
        
        if text.strip():
            # Add page separator for multi-page docs
            if page_num > 1:
                markdown_parts.append(f"\n---\n*Page {page_num}*\n")
            
            # Clean up text
            text = clean_text(text)
            markdown_parts.append(text)
    
    doc.close()
    
    result = "\n\n".join(markdown_parts)
    return post_process_markdown(result)


def convert_docx(content: bytes) -> str:
    """
    Convert Word document to Markdown.
    
    Args:
        content: DOCX file bytes
    
    Returns:
        Markdown text
    """
    doc = Document(io.BytesIO(content))
    markdown_parts = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        style_name = para.style.name.lower() if para.style else ""
        
        # Handle headings
        if "heading 1" in style_name:
            markdown_parts.append(f"# {text}")
        elif "heading 2" in style_name:
            markdown_parts.append(f"## {text}")
        elif "heading 3" in style_name:
            markdown_parts.append(f"### {text}")
        elif "heading 4" in style_name:
            markdown_parts.append(f"#### {text}")
        elif "title" in style_name:
            markdown_parts.append(f"# {text}")
        elif "list" in style_name:
            markdown_parts.append(f"- {text}")
        else:
            # Check for bold/italic runs
            formatted_text = format_runs(para.runs)
            markdown_parts.append(formatted_text)
    
    # Handle tables
    for table in doc.tables:
        table_md = convert_table(table)
        markdown_parts.append(table_md)
    
    result = "\n\n".join(markdown_parts)
    return post_process_markdown(result)


def format_runs(runs) -> str:
    """Format Word runs with bold/italic."""
    parts = []
    for run in runs:
        text = run.text
        if not text:
            continue
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        parts.append(text)
    return "".join(parts)


def convert_table(table) -> str:
    """Convert Word table to Markdown table."""
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            # Add header separator
            rows.append("|" + "|".join(["---"] * len(cells)) + "|")
    return "\n".join(rows)


def convert_html(content: str, base_url: Optional[str] = None) -> str:
    """
    Convert HTML to Markdown.
    
    Args:
        content: HTML string
        base_url: Base URL for resolving relative links
    
    Returns:
        Markdown text
    """
    # Parse and clean HTML
    soup = BeautifulSoup(content, "html.parser")
    
    # Remove script and style elements
    for element in soup(["script", "style", "nav", "footer", "header", "aside"]):
        element.decompose()
    
    # Try to find main content
    main_content = (
        soup.find("main") or 
        soup.find("article") or 
        soup.find(class_=re.compile(r"content|article|post|entry")) or
        soup.find("body") or
        soup
    )
    
    # Convert to markdown
    markdown = md(
        str(main_content),
        heading_style="atx",
        bullets="-",
        code_language="",
    )
    
    return post_process_markdown(markdown)


async def convert_url_via_markdown_new(url: str, timeout: float = 30.0) -> str:
    """
    Convert URL via markdown.new service.
    
    Args:
        url: URL to convert
        timeout: Request timeout in seconds
    
    Returns:
        Markdown text
    """
    markdown_new_url = f"{MARKDOWN_NEW_URL}/{url}"
    async with httpx.AsyncClient(timeout=timeout, follow_redirects=True) as client:
        resp = await client.get(markdown_new_url)
        resp.raise_for_status()
        
        text = resp.text
        # markdown.new returns with metadata header, extract content
        if "Markdown Content:" in text:
            # Extract content after "Markdown Content:"
            parts = text.split("Markdown Content:", 1)
            if len(parts) > 1:
                return parts[1].strip()
        return text


async def convert_url(
    url: str, 
    timeout: float = 30.0, 
    use_fallback: bool = True,
    prefer_markdown_new: bool = False
) -> str:
    """
    Fetch URL and convert to Markdown.
    
    Args:
        url: URL to fetch
        timeout: Request timeout in seconds
        use_fallback: Whether to fallback to markdown.new on failure
        prefer_markdown_new: Use markdown.new as primary source
    
    Returns:
        Markdown text
    """
    # If prefer markdown.new, use it directly
    if prefer_markdown_new:
        try:
            return await convert_url_via_markdown_new(url, timeout)
        except Exception:
            pass  # Fall through to local conversion
    
    # Try local conversion first
    try:
        async with httpx.AsyncClient(timeout=timeout, follow_redirects=True) as client:
            # Try to get markdown directly (Cloudflare Markdown for Agents)
            headers = {"Accept": "text/markdown, text/html", "Accept-Encoding": "identity"}
            resp = await client.get(url, headers=headers)
            
            content_type = resp.headers.get("content-type", "")
            
            if "text/markdown" in content_type:
                # Server returned markdown directly!
                return resp.text
            elif "application/pdf" in content_type:
                return convert_pdf(resp.content)
            else:
                # HTML - convert to markdown
                result = convert_html(resp.text, base_url=url)
                if result and len(result) > 100:  # Basic quality check
                    return result
                raise ValueError("Conversion result too short")
    except Exception as e:
        if use_fallback:
            # Fallback to markdown.new
            try:
                return await convert_url_via_markdown_new(url, timeout)
            except Exception as fallback_error:
                raise Exception(f"Both local and fallback conversion failed: {str(e)} / {str(fallback_error)}")
        raise


def clean_text(text: str) -> str:
    """Clean up extracted text."""
    # Normalize whitespace
    text = re.sub(r"[ \t]+", " ", text)
    # Remove excessive newlines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def post_process_markdown(text: str) -> str:
    """Post-process markdown for cleanup."""
    # Remove excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Clean up list formatting
    text = re.sub(r"\n\s*-\s+", "\n- ", text)
    # Remove trailing whitespace
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    return text.strip()
